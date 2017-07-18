################################################################################
# Authors: Zsombor Szira (based on original script by Jose Moreno)             #
#                                                                              #
#                                                                              #
# March 2017                                                                   #
#                                                                              #
# Takes an input json config of a tenant, that can be obtained by "Saving as"  #
#   from the APIC GUI. It produces a docx document with drawings for:          #
#     - BDs with SVIs and L3outs for each VRF                                  #
#     - Contracts and connected EPGs/L3outs                                    #
#     - supporting text, incl. EPG to interface and L3out static routes        #
#                                                                              #
################################################################################

import json
import pdb
import glob
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.shared import Pt
from docx.shared import Cm
import pydot
from six import iteritems

json_filename = '/Python/ACI/tn-EMEA-CH-SCUN-ACI-CORE.json'
docx_filename = '/Python/ACI/SCUN-ACI-Tenant-config.docx'

document = Document()
obj_styles = document.styles
obj_parstyle = obj_styles.add_style('Code Style', WD_STYLE_TYPE.PARAGRAPH)
obj_font = obj_parstyle.font
obj_font.size = Pt(11)
obj_font.name = 'Cambria'

unique_keys = set()

def _nested_lookup(key, document):
    """Lookup a key in a nested document, yield a value"""
    if isinstance(document, list):
        for d in document:
            for result in _nested_lookup(key, d):
                yield result

    if isinstance(document, dict):
        for k, v in iteritems(document):
            if k == key:
                yield v
            elif isinstance(v, dict):
                for result in _nested_lookup(key, v):
                    yield result
            elif isinstance(v, list):
                for d in v:
                    for result in _nested_lookup(key, d):
                        yield result

# Find the BDs associated to a specific context
def getBdForCtx(bds, ctxName):
    ctxbds=[]
    for bd in bds:
        for child in bd['children']:
            if 'fvRsCtx' in child:
                if child['fvRsCtx']['attributes']['tnFvCtxName'] == ctxName:
                     ctxbds.append(bd['attributes']['name'])
    return ctxbds

# Find the L3outs associated to a specific context
def getL3outsForCtx(l3outs, ctxName):
    ctxl3outs=[]
    for l3out in l3outs:
        for child in l3out['children']:
            if 'l3extRsEctx' in child:
                if child['l3extRsEctx']['attributes']['tnFvCtxName'] == ctxName:
                    ctxl3outs.append(l3out['attributes']['name'])
    return ctxl3outs

# Gets the contracts consumed by a specific L3out
def getConsContractsL3out(l3out):
    contractlist = []
    for l3outchild in l3out['children']:
        if 'l3extInstP' in l3outchild:
            for child in l3outchild['l3extInstP']['children']:
                if 'fvRsCons' in child:
                    contractlist.append(child['fvRsCons']['attributes']['tnVzBrCPName'])
    return contractlist

# Gets the subnets announced by a specific L3out
def getExportSubnets(l3out):
    subnetlist = []
    for l3outchild in l3out['children']:
        if 'l3extInstP' in l3outchild:
            for child in l3outchild['l3extInstP']['children']:
                if 'l3extSubnet' in child:
                    if child['l3extSubnet']['attributes']['scope'] == 'export-rtctrl':
                        subnetlist.append(child['l3extSubnet']['attributes']['ip'])
    return subnetlist

# Gets the subnets imported by a specific L3out
def getImportSubnets(l3out):
    subnetlist = []
    for l3outchild in l3out['children']:
        if 'l3extInstP' in l3outchild:
            for child in l3outchild['l3extInstP']['children']:
                if 'l3extSubnet' in child:
                    if child['l3extSubnet']['attributes']['scope'] == 'import-security':
                        subnetlist.append(child['l3extSubnet']['attributes']['ip'])
    return subnetlist

# Get contracts consumed by an EPG
def getConsContractsforEPG (epg):
    contractlist = []
    for child in epg['children']:
        if 'fvRsCons' in child:
            contractlist.append(child['fvRsCons']['attributes']['tnVzBrCPName'])
    return contractlist

# Get contracts provided by an EPG
def getProvContractsforEPG (epg):
    contractlist = []
    for child in epg['children']:
        if 'fvRsProv' in child:
            contractlist.append(child['fvRsProv']['attributes']['tnVzBrCPName'])
    return contractlist

# Gets the BD name in an EPG
def getBdForEPG (epg):
    for child in epg['children']:
        if 'fvRsBd' in child:
            return child['fvRsBd']['attributes']['tnFvBDName']

# Gets the IP addresses (subnets) defined in a BD
def getIpsForBD (bds, bd_name):
    ips=[]
    for bd in bds:
        if bd['attributes']['name'] == bd_name:
            for child in bd['children']:
                if 'fvSubnet' in child:
                    ips.append(child['fvSubnet']['attributes']['ip'])
    return ips

#Get EPGs consuming a contract
def getEPGForConsumeCont (contcons):
    consepg=[]
    for epg in fvAllEPG:
        for child in epg['children']:
            if 'fvRsCons' in child:
                if child['fvRsCons']['attributes']['tnVzBrCPName'] == contcons:
                    consepg.append(epg['attributes']['name'])
    return consepg

#Get EPGs providing a contract
def getEPGForProvideCont (contcons):
    provepg=[]
    for epg in fvAllEPG:
        for child in epg['children']:
            if 'fvRsProv' in child:
                if child['fvRsProv']['attributes']['tnVzBrCPName'] == contcons:
                    provepg.append(epg['attributes']['name'] + ' - ' + epg['attributes']['descr'])
    return provepg

#Get EPGs Providing or Consuming a contract
def getEPGForCont (contcons):
    usebyepg=[]
    for epg in fvAllEPG:
        for child in epg['children']:
            if 'fvRsProv' in child and child['fvRsProv']['attributes']['tnVzBrCPName'] == contcons:
                 usebyepg.append(epg['attributes']['name'])
            else:
                if 'fvRsCons' in child and child['fvRsCons']['attributes']['tnVzBrCPName'] == contcons:
                    usebyepg.append(epg['attributes']['name'])
    return list(set(usebyepg))

#Get L3outs Providing or Consuming a contract
def getL3outForCont (contcons):
    usebyl3out=[]
    for l3out in l3extOut:
        for l3outchild in l3out['children']:
            if 'l3extInstP' in l3outchild:
                for child in l3outchild['l3extInstP']['children']:
                    if 'fvRsCons' in child and child['fvRsCons']['attributes']['tnVzBrCPName'] == contcons:
                        usebyl3out.append(l3out['attributes']['name'])
                    else:
                        if 'fvRsProv' in child and child['fvRsProv']['attributes']['tnVzBrCPName'] == contcons:
                            usebyl3out.append(l3out['attributes']['name'])
    return list(set(usebyl3out))

def getL3outStatic(l3outname):
    temp = []
    for child in l3extOut:
        if child['attributes']['name'] == l3outname:
            for nodep in child['children']:
                if 'l3extLNodeP' in nodep:
                    temp.append(nodep['l3extLNodeP'])
    statics = []
    statics = _nested_lookup('ip',temp)
    return list(set(statics))

# Get path for an EPG
def getPathEPG (epg):
    pathlist = []
    encaplist = []
    modelist = []
    for child in epg['children']:
        if 'fvRsPathAtt' in child:
            pathlist.append(child['fvRsPathAtt']['attributes']['tDn'])
            encaplist.append(child['fvRsPathAtt']['attributes']['encap'])
            modelist.append(child['fvRsPathAtt']['attributes']['mode'])
    matrix = []
    for i in range(len(pathlist)):
        matrix.append([pathlist[i] + '; ', encaplist[i] + '; ', modelist[i]])
    return matrix

# Load the config file in a dictionary
with open(json_filename) as fd:
    config = json.load(fd)
    fd.close()

# Strip off the initial metadata labels, find out the tenant name
fvTenant = config['imdata'][0]['fvTenant']
document.add_heading('ACI Configuration for tenant ' + fvTenant['attributes']['name'], 0)

# Some lists with JSON code, to make things easier
fvAp = []
fvCtx = []
l3extOut = []
fvBD = []
vzBrCP = []
fvAllEPG = []


# Get the list of ANPs, this returns a list so I need to go trough it again.
# To create the document in order I need to divide the objects depending on their class
for child in fvTenant['children']:
    if 'fvAp' in child:
        fvAp.append(child['fvAp'])
    if 'fvCtx' in child:
        fvCtx.append(child['fvCtx'])
    if 'l3extOut' in child:
        l3extOut.append(child['l3extOut'])
    if 'fvBD' in child:
        fvBD.append(child['fvBD'])
    if 'vzBrCP' in child:
        vzBrCP.append(child['vzBrCP'])
for anp in fvAp:
    for child in anp['children']:
         if 'fvAEPg' in child:
            fvAllEPG.append(child['fvAEPg'])


# Networking info (VRFs, BDs, L3outs)
document.add_heading('Networking configuration', level=1)

# Print a brief description of the private networks
document.add_paragraph('', style='Code Style')
document.add_heading('Private networks (VRFs)', level=2)
document.add_paragraph('', style='Code Style')
if fvCtx.__len__() == 0:
    document.add_paragraph('This tenant has no private networks (VRFs) defined, it is probably using a private network defined in the common tenant.', style='Code Style')
else:
    document.add_paragraph('This tenant has the following private networks (VRFs) defined:', style='Code Style')
    for vrf in fvCtx:
        paragraph=vrf['attributes']['name']
        document.add_heading(paragraph, level=3)
        # Bridge domains
        ctxbds = getBdForCtx(fvBD,vrf['attributes']['name'])
        if ctxbds.__len__() == 0:
            paragraph = 'No bridge domains associated to this VRF.'
            document.add_paragraph(paragraph, style='Code Style')
        else:
            paragraph = 'The following bridge domains are defined in this VRF: '
            document.add_paragraph(paragraph, style='Code Style')
            for bd in ctxbds:
                paragraph = bd
                ips = getIpsForBD(fvBD, bd)
                if ips.__len__() == 0:
                    paragraph = paragraph + '. subnets: N/A'
                else:
                    paragraph = paragraph + '. subnets: '
                    for ip in ips:
                        paragraph = paragraph + ip + ", "
                    # Remove last comma
                    paragraph = paragraph[0:paragraph.__len__()-2]
                document.add_paragraph(paragraph, style='List Bullet')
        # L3 outs
        l3outs = getL3outsForCtx(l3extOut,vrf['attributes']['name'])
        if l3outs.__len__() == 0:
            paragraph = 'No external IP connections associated to this VRF'
            document.add_paragraph(paragraph, style='Code Style')
        else:
            paragraph = 'The following external IP connections are associated to his VRF: '
            document.add_paragraph(paragraph, style='Code Style')
            for l3out in l3outs:
                paragraph = l3out
                statroute = []
                statroute = getL3outStatic(l3out)
                if statroute.__len__() == 0:
                    paragraph = paragraph + '. static routes: N/A'
                else:
                    paragraph = paragraph + '. subnets: '
                    for stat in statroute:
                        paragraph = paragraph + stat + ', '
                document.add_paragraph(paragraph, style='List Bullet')

        # Create image
        graph = pydot.Dot(graph_type='graph')
        for bd in ctxbds:
            edge = pydot.Edge('VRF '+ vrf['attributes']['name'], 'BD ' + bd)
            graph.add_edge(edge)
            ips = getIpsForBD(fvBD, bd)
            for ip in ips:
                edge = pydot.Edge('BD ' + bd, ip)
                graph.add_edge(edge)
        for l3out in l3outs:
            edge = pydot.Edge('VRF ' + vrf['attributes']['name'], 'L3out ' + l3out)
            graph.add_edge(edge)
        edgelist = []
        edgelist = graph.get_edge_list()
        if not edgelist.__len__() == 0:
            graph.write_png(vrf['attributes']['name'] + '.png')
            document.add_picture(vrf['attributes']['name'] + '.png', width=Cm(18.0))

# List all EPGs with BD and associated interfaces
document.add_paragraph(' ', style='Code Style')
document.add_heading('EPGs', level=1)
document.add_paragraph(' ', style='Code Style')
if fvAllEPG.__len__() == 0:
    document.add_paragraph('This tenant has no EPGs defined.', style='Code Style')
else:
    document.add_paragraph('The following EPGs are defined in this tenant:', style='Code Style')
    for epgs in fvAllEPG:
        document.add_paragraph('', style='Code Style')
        document.add_heading('EPG ' + epgs['attributes']['name'], level=2)
        bd = getBdForEPG(epgs)
        document.add_paragraph('Bridge Domain: ' + bd)
        document.add_paragraph('Path:')
        pathmatrix = getPathEPG(epgs)
        for i in range(len(pathmatrix)):
            document.add_paragraph(pathmatrix[i], style='List Bullet')

# Print one section for each Contract
document.add_paragraph(' ', style='Code Style')
document.add_heading('Contracts', level=1)
document.add_paragraph(' ', style='Code Style')
if vzBrCP.__len__() == 0:
    document.add_paragraph('This tenant has no contracts defined.', style='Code Style')
else:
    document.add_paragraph('The following contracts are defined in this tenant:', style='Code Style')
    for cont in vzBrCP:
        document.add_paragraph('', style='Code Style')
        document.add_heading('Contract ' + cont['attributes']['name'], level=2)
        UseEPG = []
        UseEPG = getEPGForCont(cont['attributes']['name'])
        UseL3Out = []
        UseL3Out = getL3outForCont(cont['attributes']['name'])
        if UseEPG.__len__() == 0:
            document.add_paragraph('This contract is not used by EPGs.', style='Code Style')
        else:
            document.add_paragraph('The following EPGs Use this Contract:', style='Code Style')
            for epg in UseEPG:
                document.add_paragraph('EPG ' + epg, style='List Bullet')
        if UseL3Out.__len__() == 0:
            document.add_paragraph('This contract is not used by L3 Outs.', style='Code Style')
        else:
            document.add_paragraph('The following L3 Outs use this Contract:', style='Code Style')
            for l3out in UseL3Out:
                document.add_paragraph('L3 Out ' + l3out, style='List Bullet')
        # Create image
        graph = pydot.Dot(graph_type='digraph')
        node_cont = pydot.Node('Contract ' + cont['attributes']['name'])
        graph.add_node(node_cont)
        if UseEPG.__len__() != 0:
            subg_epg = pydot.Subgraph(rank='max')
            for epg in UseEPG:
                node_epg = pydot.Node('EPG\n ' + epg)
                subg_epg.add_node(node_epg)
                graph.add_edge(pydot.Edge(node_cont, node_epg, dir='both'))
            graph.add_subgraph(subg_epg)
        if UseL3Out.__len__() != 0:
            subg_l3out = pydot.Subgraph(rank='min')
            for l3out in UseL3Out:
                node_l3out = pydot.Node('L3 out\n ' + l3out)
                subg_l3out.add_node(node_l3out)
                graph.add_edge(pydot.Edge(node_cont, node_l3out, dir='both'))
            graph.add_subgraph(subg_l3out)
        graph.write_png(cont['attributes']['name'] + '.png')
        document.add_picture(cont['attributes']['name'] + '.png', width=Cm(18.0))

print ("Saving document...")
document.save(docx_filename)
